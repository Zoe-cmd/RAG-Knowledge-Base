"""
Word (docx) 文档解析器。

使用 python-docx 库解析 docx 文档，提取段落与表格文本。
忽略图片、公式等非文本内容。
"""

from pathlib import Path

from app.parsers.base import DocumentParser


class DocxParser(DocumentParser):
    """Word 文档解析器。

    提取 docx 中的段落与表格文本。
    表格按行转换为文本，单元格用制表符分隔。
    """

    def parse(self, file_path: Path) -> str:
        """解析 docx 为纯文本。

        Args:
            file_path: docx 文件路径

        Returns:
            解析后的纯文本（段落 + 表格）

        Raises:
            Exception: docx 解析失败
        """
        import docx

        doc = docx.Document(str(file_path))
        text_parts = []

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 提取表格（按行转换为文本）
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells
                )
                if row_text.strip():
                    text_parts.append(row_text)

        raw_text = "\n\n".join(text_parts)
        return self.clean_text(raw_text)
