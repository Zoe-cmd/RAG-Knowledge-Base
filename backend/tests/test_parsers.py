"""
文档解析器单元测试。

测试 PDF、DOCX、Markdown、TXT 解析器与工厂。
使用临时文件测试各解析器的解析逻辑。
"""

import tempfile
from pathlib import Path

import pytest

from app.parsers.docx_parser import DocxParser
from app.parsers.factory import get_parser, parse_document
from app.parsers.markdown_parser import MarkdownParser
from app.parsers.pdf_parser import PDFParser, ScannedPDFError
from app.parsers.txt_parser import TxtParser


class TestTxtParser:
    """TXT 解析器测试。"""

    def test_utf8_text(self, tmp_path):
        """UTF-8 编码文本。"""
        file_path = tmp_path / "test.txt"
        file_path.write_text("这是测试文本。\nHello World.", encoding="utf-8")
        parser = TxtParser()
        result = parser.parse(file_path)
        assert "这是测试文本" in result
        assert "Hello World" in result

    def test_gbk_text(self, tmp_path):
        """GBK 编码文本。"""
        file_path = tmp_path / "test.txt"
        file_path.write_text("这是GBK编码的文本。", encoding="gbk")
        parser = TxtParser()
        result = parser.parse(file_path)
        assert "GBK" in result or "这是" in result

    def test_empty_file(self, tmp_path):
        """空文件。"""
        file_path = tmp_path / "empty.txt"
        file_path.write_text("", encoding="utf-8")
        parser = TxtParser()
        result = parser.parse(file_path)
        assert result == ""


class TestMarkdownParser:
    """Markdown 解析器测试。"""

    def test_basic_markdown(self, tmp_path):
        """基本 Markdown 解析。"""
        file_path = tmp_path / "test.md"
        file_path.write_text(
            "# 标题\n\n这是正文内容。\n\n## 子标题\n\n- 列表项1\n- 列表项2",
            encoding="utf-8",
        )
        parser = MarkdownParser()
        result = parser.parse(file_path)
        # 应去除 Markdown 语法标记，保留文本
        assert "标题" in result
        assert "正文内容" in result
        assert "列表项1" in result

    def test_markdown_links(self, tmp_path):
        """Markdown 链接解析。"""
        file_path = tmp_path / "test.md"
        file_path.write_text(
            "[OpenAI 官网](https://openai.com)", encoding="utf-8"
        )
        parser = MarkdownParser()
        result = parser.parse(file_path)
        # 链接应被移除，保留链接文本
        assert "OpenAI 官网" in result
        assert "https://openai.com" not in result

    def test_markdown_code_block(self, tmp_path):
        """Markdown 代码块解析。"""
        file_path = tmp_path / "test.md"
        file_path.write_text(
            "```python\nprint('hello')\n```", encoding="utf-8"
        )
        parser = MarkdownParser()
        result = parser.parse(file_path)
        assert "print" in result


class TestDocxParser:
    """DOCX 解析器测试。"""

    def test_parse_docx(self, tmp_path):
        """解析 DOCX 文件（需 python-docx）。"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        file_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("这是第一段。")
        doc.add_paragraph("这是第二段。")
        doc.save(str(file_path))

        parser = DocxParser()
        result = parser.parse(file_path)
        assert "第一段" in result
        assert "第二段" in result


class TestPDFParser:
    """PDF 解析器测试。"""

    def test_scanned_pdf_raises_error(self, tmp_path):
        """扫描件 PDF 应抛出 ScannedPDFError。"""
        # 创建一个空的 PDF（无文本内容）
        try:
            from PyPDF2 import PdfWriter

            file_path = tmp_path / "empty.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=200, height=200)
            writer.write(str(file_path))

            parser = PDFParser()
            with pytest.raises(ScannedPDFError):
                parser.parse(file_path)
        except ImportError:
            pytest.skip("PyPDF2 未安装")


class TestParserFactory:
    """解析器工厂测试。"""

    def test_get_parser_pdf(self):
        """获取 PDF 解析器。"""
        parser = get_parser("pdf")
        assert isinstance(parser, PDFParser)

    def test_get_parser_docx(self):
        """获取 DOCX 解析器。"""
        parser = get_parser("docx")
        assert isinstance(parser, DocxParser)

    def test_get_parser_md(self):
        """获取 Markdown 解析器。"""
        parser = get_parser("md")
        assert isinstance(parser, MarkdownParser)

    def test_get_parser_txt(self):
        """获取 TXT 解析器。"""
        parser = get_parser("txt")
        assert isinstance(parser, TxtParser)

    def test_get_parser_case_insensitive(self):
        """文件类型应大小写不敏感。"""
        parser = get_parser("PDF")
        assert isinstance(parser, PDFParser)

    def test_get_parser_unsupported(self):
        """不支持的文件类型应抛出 ValueError。"""
        with pytest.raises(ValueError, match="不支持"):
            get_parser("xlsx")

    def test_get_parser_singleton(self):
        """相同类型应返回同一实例（lru_cache）。"""
        p1 = get_parser("txt")
        p2 = get_parser("txt")
        assert p1 is p2

    def test_parse_document_convenience(self, tmp_path):
        """parse_document 便捷函数。"""
        file_path = tmp_path / "test.txt"
        file_path.write_text("测试文本", encoding="utf-8")
        result = parse_document(file_path, "txt")
        assert "测试文本" in result


class TestCleanText:
    """文本清洗测试。"""

    def test_clean_text_merges_spaces(self):
        """合并连续空格。"""
        from app.parsers.base import DocumentParser

        result = DocumentParser.clean_text("a  b   c")
        assert result == "a b c"

    def test_clean_text_merges_newlines(self):
        """合并多个换行。"""
        from app.parsers.base import DocumentParser

        result = DocumentParser.clean_text("a\n\n\n\n\nb")
        assert "\n\n\n" not in result

    def test_clean_text_strips(self):
        """去除首尾空白。"""
        from app.parsers.base import DocumentParser

        result = DocumentParser.clean_text("  hello  ")
        assert result == "hello"

    def test_clean_text_empty(self):
        """空文本。"""
        from app.parsers.base import DocumentParser

        assert DocumentParser.clean_text("") == ""
        assert DocumentParser.clean_text(None) == ""
